import streamlit as st
import zipfile
import os
import tempfile
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image

# ---------------------------
# 1. Interface upload
# ---------------------------
st.title("📑 Générateur de rapport vétérinaire")
pdf_file = st.file_uploader("Uploader le rapport PDF", type=["pdf"])
zip_file = st.file_uploader("Uploader le fichier ZIP d’images", type=["zip"])

if pdf_file and zip_file:
    # ---------------------------
    # 2. Extraction du texte du PDF
    # ---------------------------
    pdf_reader = PdfReader(pdf_file)
    pdf_text = ""
    for page in pdf_reader.pages:
        pdf_text += page.extract_text() + "\n"
    st.subheader("📝 Texte extrait du rapport")
    st.text_area("Texte du rapport", pdf_text, height=200)

    # ---------------------------
    # 3. Extraction des images du ZIP
    # ---------------------------
    temp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    images = [os.path.join(temp_dir, f) for f in os.listdir(temp_dir) 
              if f.lower().endswith((".png", ".jpg", ".jpeg"))]

    selected_images = []
    st.subheader("🖼️ Sélectionner les images clés")
    for img_path in images:
        img = Image.open(img_path)
        st.image(img, caption=os.path.basename(img_path), width=250)
        if st.checkbox(f"Sélectionner {os.path.basename(img_path)}"):
            selected_images.append(img_path)

    # ---------------------------
    # 4. Génération du rapport final
    # ---------------------------
    if st.button("📤 Générer le rapport médical"):
        doc = Document()
        doc.add_heading("Rapport Médical", level=1)
        doc.add_paragraph(pdf_text)

        doc.add_heading("Images Clés", level=2)
        for img_path in selected_images:
            doc.add_picture(img_path, width=docx.shared.Inches(3))
            doc.add_paragraph(os.path.basename(img_path))

        output_path = os.path.join(temp_dir, "rapport_final.docx")
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button("⬇️ Télécharger le rapport final", f, file_name="rapport_final.docx")
